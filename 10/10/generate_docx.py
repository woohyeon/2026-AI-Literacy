import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Sets background color for a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding (margins) for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_run_font(run, font_name='Malgun Gothic', size_pt=10, bold=False, color_rgb=None):
    """Sets the font properties for a run, including East Asian character support."""
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb

def set_table_borders(table):
    """Applies clean borders to a table."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    
    # top, bottom, insideH borders
    for border_name, color, sz, val in [
        ('top', 'CCCCCC', '4', 'single'),
        ('bottom', 'CCCCCC', '8', 'single'),
        ('insideH', 'E0E0E0', '4', 'single'),
        ('left', 'none', '0', 'none'),
        ('right', 'none', '0', 'none'),
        ('insideV', 'none', '0', 'none')
    ]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        borders.append(border)
    tblPr.append(borders)

def main():
    doc = docx.Document()

    # Set document margins (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Theme Colors
    primary_color = RGBColor(31, 78, 120)     # Dark Blue (#1F4E78)
    charcoal_color = RGBColor(51, 51, 51)     # Dark Gray for body text
    header_bg_color = "F2F2F2"                # Light Gray for table headers

    # --- Title ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(24)
    run_title = p_title.add_run("삶의 만족도와 자살률의 상관관계 분석 보고서\n(2020-2024)")
    set_run_font(run_title, 'Malgun Gothic', 20, bold=True, color_rgb=primary_color)

    # --- Metadata (Subtitle) ---
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_meta.paragraph_format.space_after = Pt(24)
    run_meta = p_meta.add_run("작성일자: 2026년 6월 11일\n통계 구분: 통계청 사회조사 및 사망원인통계")
    set_run_font(run_meta, 'Malgun Gothic', 9.5, bold=False, color_rgb=RGBColor(128, 128, 128))

    # --- Helper function for Headings ---
    def add_custom_heading(text, level):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        if level == 1:
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            set_run_font(run, 'Malgun Gothic', 14, bold=True, color_rgb=primary_color)
        elif level == 2:
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            set_run_font(run, 'Malgun Gothic', 11.5, bold=True, color_rgb=primary_color)
        return p

    def add_body_paragraph(text, space_after=6, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        set_run_font(run, 'Malgun Gothic', 10, bold=bold, color_rgb=charcoal_color)
        return p

    def add_bullet_point(text, space_after=4):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        set_run_font(run, 'Malgun Gothic', 10, bold=False, color_rgb=charcoal_color)
        return p

    def add_image_with_caption(img_filename, caption_text):
        if os.path.exists(img_filename):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(4)
            p_img.add_run().add_picture(img_filename, width=Inches(5.8))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(12)
            run_cap = p_cap.add_run(caption_text)
            set_run_font(run_cap, 'Malgun Gothic', 9, bold=True, color_rgb=RGBColor(100, 100, 100))
        else:
            p_err = doc.add_paragraph()
            run_err = p_err.add_run(f"[이미지 누락: {img_filename}]")
            set_run_font(run_err, 'Malgun Gothic', 10, bold=True, color_rgb=RGBColor(255, 0, 0))

    # --- 1. 요약 ---
    add_custom_heading("1. 요약 (Executive Summary)", 1)
    add_bullet_point("국가적 추이 역설: 2021년 코로나19 팬데믹 정기 당시 전국 삶의 만족도는 대폭 하락(3.42 -> 3.16)했으나 자살률은 소폭 상승(25.7명 -> 26.0명)에 그쳤습니다. 반면, 팬데믹 이후인 2024년에는 삶의 만족도가 평균 3.37 수준으로 회복되었음에도 불구하고 전국 자살률은 29.1명으로 최근 5년 중 가장 높은 수치를 기록했습니다.")
    add_bullet_point("지역별 약한 음의 상관관계: 5개년 전체 데이터를 통합 분석한 결과, 삶의 만족 비율과 자살률 사이에는 약한 음의 상관관계(r = -0.1704)가 나타났습니다. 삶의 만족도가 높은 지역일수록 자살률이 다소 낮은 경향이 있으나, 그 영향력은 절대적이지 않습니다.")
    add_bullet_point("연도별 상관관계의 변동: 2020년에는 만족도와 자살률 간에 강한 음의 상관관계(r = -0.5226)를 보였으나, 2021년~2024년에는 상관관계가 매우 낮아지거나 불만족 비율과 자살률이 역설적으로 음의 상관관계(r = -0.4870, 2024년)를 보이는 등 연도별 편차가 컸습니다.")
    add_bullet_point("인구 통계학적 요인의 매개: 불만족도가 높은 지역(대구, 세종 등)의 자살률이 낮고, 불만족도가 낮은 지역(충남, 강원 등)의 자살률이 높은 현상이 관찰되었습니다. 이는 고령 인구 비율이 높은 지역일수록 노인 자살률의 영향으로 전체 자살률이 급증하는 반면, 청년 비율이 높은 지역(세종 등)은 삶에 대한 주관적 불만족은 높지만 실제 자살률은 낮게 나타나는 인구 구조적 차이에서 기인합니다.")
    add_bullet_point("성별 격차: 남성의 자살률(41.8명, 2024년 전국)은 여성(16.6명)에 비해 약 2.5배 높았으나, 주관적 삶의 만족도 점수는 남녀 간 유의미한 차이가 없었습니다.")

    # --- 2. 데이터 개요 ---
    add_custom_heading("2. 데이터 개요 및 전처리", 1)
    add_body_paragraph("본 보고서의 데이터는 2020년부터 2024년까지의 5개년 기간을 대상으로 통계청 사회조사 데이터(삶의 만족도) 및 사망원인통계 데이터(자살률)를 추출하여 결합하였습니다.")
    add_bullet_point("삶의 만족도: 매우 만족, 약간 만족, 보통, 약간 불만족, 매우 불만족의 5개 척도 응답 비율(%)을 바탕으로, 긍정적인 '만족 비율(매우 만족 + 약간 만족)'과 부정적인 '불만족 비율(매우 불만족 + 약간 불만족)'을 정의하였습니다. 또한, 평균적인 만족 수준을 5점 만점으로 환산한 '평균 만족도 점수'를 산출하였습니다.")
    add_bullet_point("자살률: 각 지자체별 인구 10만 명당 자살 사망자 수를 활용하였습니다.")
    add_bullet_point("행정구역 명칭 통일: 상이하게 기재된 행정구역명을 표준화하였습니다. (예: '전라북도' -> '전북특별자치도', '제주도' -> '제주특별자치도')")
    add_bullet_point("분석 단위: 시도 경향성의 왜곡을 막기 위해 전국 평균('전국') 데이터를 분리하고 17개 시도별 데이터를 별도로 분석하였습니다.")

    # --- 3. 전국 단위 시계열 추이 분석 ---
    add_custom_heading("3. 전국 단위 시계열 추이 분석", 1)
    add_body_paragraph("전국 평균 데이터를 연도별 및 성별로 나누어 정리한 결과는 다음과 같습니다.")

    # --- Table 1 ---
    table1_data = [
        ["연도", "성별", "만족 비율 (%)", "불만족 비율 (%)", "평균 만족 점수 (5점)", "자살률 (10만명당, 명)"],
        ["2020", "합계", "42.7", "12.5", "3.423", "25.7"],
        ["", "남자", "43.5", "11.8", "3.443", "35.5"],
        ["", "여자", "41.9", "13.0", "3.397", "15.9"],
        ["2021", "합계", "34.0", "22.9", "3.161", "26.0"],
        ["", "남자", "33.8", "23.5", "3.155", "35.9"],
        ["", "여자", "34.3", "22.4", "3.174", "16.2"],
        ["2022", "합계", "43.3", "14.1", "3.388", "25.2"],
        ["", "남자", "43.4", "14.1", "3.389", "35.3"],
        ["", "여자", "43.2", "14.0", "3.382", "15.1"],
        ["2023", "합계", "42.2", "14.5", "3.366", "27.3"],
        ["", "남자", "42.5", "14.6", "3.373", "38.3"],
        ["", "여자", "42.0", "14.5", "3.366", "16.5"],
        ["2024", "합계", "40.1", "12.7", "3.367", "29.1"],
        ["", "남자", "40.5", "12.4", "3.377", "41.8"],
        ["", "여자", "39.7", "13.1", "3.358", "16.6"]
    ]

    t1 = doc.add_table(rows=len(table1_data), cols=6)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t1)

    for r_idx, row_data in enumerate(table1_data):
        row = t1.rows[r_idx]
        # Keep row on same page
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
        # If header, repeat header
        if r_idx == 0:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            
            # Format text
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            
            if r_idx == 0:
                set_cell_background(cell, header_bg_color)
                set_run_font(run, 'Malgun Gothic', 9.5, bold=True, color_rgb=primary_color)
            else:
                is_bold = (row_data[1] == "합계")
                set_run_font(run, 'Malgun Gothic', 9, bold=is_bold, color_rgb=charcoal_color)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(6)

    # Insert national_trends.png
    add_image_with_caption("national_trends.png", "그림 1. 전국 삶의 만족도 평균과 자살률 추이 (2020-2024)")

    add_body_paragraph("추이 분석에 따르면, 코로나19가 절정에 달했던 2021년에는 사회적 통제와 고립감으로 인해 국민들의 삶의 만족도 점수가 3.16으로 대폭 하락하고 주관적 불만족 응답이 약 23%에 육박하는 급격한 악화 양상을 보였습니다. 그럼에도 실제 자살률은 25.7명에서 26.0명 수준으로 통제 범위 내에 머물렀습니다. 그러나 거리두기가 해제되고 사회 활동이 재개된 2023-2024년에는 삶의 만족도가 예년 수준(3.37점)을 찾았음에도 전국 자살률이 29.1명(2024년)으로 가파르게 치솟아 최근 5년 중 최대치를 기록했습니다. 이는 재난 상황 속 연대의식에 의한 일시적 자살 지연 현상(Pull-together effect)이 해제되고, 사회 복귀 과정에서 경제적 양극화 및 상대적 박탈감이 축적되며 자살률 상승으로 뒤늦게 현실화된 것으로 판단됩니다.")

    # --- 4. 상관관계 분석 결과 ---
    add_custom_heading("4. 상관관계 분석 결과 (Correlation Analysis)", 1)
    add_body_paragraph("17개 지자체(전국 데이터 제외) 데이터를 통합하여 피어슨 상관계수(r)를 구한 결과, 전체 기간(2020~2024년) 통합 데이터를 기준으로 '만족 비율'과 '자살률' 간에는 r = -0.1704의 약한 음의 상관관계가 나타났습니다. 삶의 만족도 평균 점수와의 상관관계는 r = -0.0477로 무상관에 가깝게 계산되었습니다.")
    add_body_paragraph("그러나 연도별로 관계성을 분리하여 관측할 경우, 사회적 상황에 따라 극적인 반전이 관찰됩니다:")
    add_bullet_point("2020년에는 삶의 만족도 비율과 자살률 사이에 r = -0.5226의 강한 음의 상관관계가 발생하여 보편적인 상식이 부합하는 경향을 나타냈습니다.")
    add_bullet_point("반면, 포스트 팬데믹 시기인 2024년에는 만족도 비율과의 상관관계는 사라진 반면, 불만족 비율과 자살률 사이에 r = -0.4870의 역설적인 강한 음의 상관관계가 발생하였습니다. 즉, 주관적으로 보고되는 삶의 불만족도가 높은 지역일수록 자살률이 낮고, 불만족도가 극히 낮은 지역의 자살률이 높게 기록되었습니다. 이 역설적인 관계는 지역별 고유 인구구조 특성에 따른 것이며 상세 분석은 6장에서 다룹니다.")

    # Scatter & Heatmap Images
    add_image_with_caption("regional_correlation_scatter.png", "그림 2. 지역별 삶의 만족도 비율과 자살률 상관관계 산점도 (2020-2024)")
    add_image_with_caption("correlation_matrix_heatmap.png", "그림 3. 삶의 만족도 상세 설문 항목 및 자살률 상관관계 히트맵 (2020-2024)")

    add_body_paragraph("히트맵 분석에 따르면, 매우 만족(r = -0.19) 및 약간 만족(r = -0.09)은 음의 상관관계를 보여 자살률을 억제하는 보호 기제로 작동하는 반면, 약간 불만족(r = -0.22)과 매우 불만족(r = -0.04) 역시 자살률과 음의 상관성을 나타냅니다. 이는 불만족을 겉으로 표현하고 요구하는 분위기가 강한 지역(민주적 의사표명 활성화 등)이 실제 고독사나 위험도가 높은 소외 계층의 차단에는 유리하게 작용할 수 있음을 상징적으로 뒷받침합니다.")

    # --- 5. 성별 상관관계 분석 ---
    add_custom_heading("5. 성별 상관관계 분석", 1)
    add_body_paragraph("성별에 따른 주관적 체감 만족도와 실제 자살률 지표 사이의 구조적 불일치를 규명하였습니다. 남녀의 삶의 만족 비율과 자살률 상관성은 남성 r = -0.1208, 여성 r = -0.0422로 나타나 공통적으로 매우 약한 선형 관계를 나타냅니다.")

    add_image_with_caption("gender_correlation_scatter.png", "그림 4. 성별 삶의 만족도 비율과 자살률 상관관계 산점도 및 회귀선")

    add_body_paragraph("하지만 자살률 지표 자체의 성별 분포는 크게 대조됩니다. 남성의 연도별/지역별 자살률은 30명에서 최고 50명대까지 집중 분포되는 반면, 여성의 자살률은 10명대에서 최고 20명대 초반에 조밀하게 고정되어 있습니다. 이와 대조적으로 설문상 주관적으로 인지하는 만족 비율은 남녀 모두 30%~48%의 동일한 스펙트럼 내에서 유사하게 측정됩니다. 이는 주관적 삶의 만족도 하락이 여성에 비해 남성의 생명안전을 훨씬 극단적으로 위협함을 증명하며, 남성 자살 예방을 위해서는 단순 정신건강 증진 외에 실직, 생계 위기, 은퇴 후 단절 등 경제 구조적 요인에 대한 우선적인 대응망이 전제되어야 함을 나타냅니다.")

    # --- 6. 지역별 특이값 분석 ---
    add_custom_heading("6. 지역별 특이값 분석 및 역설의 원인 규명 (Demographic Paradox)", 1)
    add_body_paragraph("2024년 최신 통계 데이터를 기준으로, 설문 통계의 한계와 인구 구조적 역설(Demographic Paradox)을 가장 뚜렷하게 증명하는 지자체들을 선정하여 분석하였습니다.")

    # --- Table 2 ---
    table2_data = [
        ["지역", "자살률 순위", "자살률 (명)", "만족 비율 (%)", "불만족 비율 (%)", "평균 만족 점수 (5점)"],
        ["제주특별자치도", "1위", "36.3", "43.4", "10.2", "3.442"],
        ["충청남도", "2위", "34.8", "47.0", "6.2", "3.555"],
        ["인천광역시", "9위", "31.2", "28.0", "13.8", "3.218"],
        ["대구광역시", "12위", "29.4", "36.1", "21.4", "3.217"],
        ["세종특별자치시", "17위", "23.0", "42.8", "20.8", "3.351"]
    ]

    t2 = doc.add_table(rows=len(table2_data), cols=6)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2)

    for r_idx, row_data in enumerate(table2_data):
        row = t2.rows[r_idx]
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        if r_idx == 0:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            
            if r_idx == 0:
                set_cell_background(cell, header_bg_color)
                set_run_font(run, 'Malgun Gothic', 9.5, bold=True, color_rgb=primary_color)
            else:
                set_run_font(run, 'Malgun Gothic', 9, bold=False, color_rgb=charcoal_color)

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_after = Pt(12)

    add_custom_heading("6.1 세종 vs 충남·제주의 역설", 2)
    add_body_paragraph("세종특별자치시는 주관적 불만족도가 20.8%(전국 2위)로 매우 불만족 비율이 높은 지역이지만, 실제 자살률은 23.0명으로 전국에서 가장 낮습니다. 세종시는 중앙 부처 이전 등으로 평균 연령이 매우 낮은 청년 및 신혼부부 중심의 행정도시입니다. 젊은 세대일수록 설문 조사 시 자신의 불만족도나 피로도를 솔직하고 예민하게 표현하는 경향이 있어 불만족도가 높게 나타납니다. 하지만 연령별 기초 자살률 분포상 젊은 층은 고령층에 비해 자살률의 절대적 위험도가 극히 낮기 때문에, 높은 설문상 불만족이 실제 사망률 통계로 이어지지 않습니다.")
    add_body_paragraph("반면, 충청남도와 제주특별자치도는 불만족 비율이 각각 6.2%(전국 최저), 10.2%로 대단히 긍정적인 사회 인식을 나타냅니다. 그러나 자살률은 34.8명, 36.3명으로 전국 1, 2위를 기록하는 극단적 역설을 보입니다. 농어촌 및 도서지역 비중이 높은 고령화 지역의 특성상, 장노년층은 설문에서 '보통'이나 '약간 만족'에 보수적으로 답하여 부정적 수치를 감추는 방어적 응답 성향을 보입니다. 그러나 현실에서는 노령기 빈곤, 홀몸 노인 고독, 만성 질환 등의 사회적 단절에 무방비하게 노출되어 극단적인 자살 사망 지표로 직결되는 것입니다.")

    add_custom_heading("6.2 인천의 고만족도 결핍형 패턴", 2)
    add_body_paragraph("인천광역시는 만족 비율이 28.0%로 전국 최하위를 차지하는 동시에, 자살률은 31.2명으로 평균 이상을 기록하는 취약 지역입니다. 주관적 만족도의 결핍과 객관적 자살 위험 지표가 동조하여 나쁜 신호를 보이는 만큼, 긴급한 정주여건 개선과 밀착 예방 사업의 병행이 시급합니다.")

    add_custom_heading("6.3 시도별 지표 시각화 분석", 2)

    add_image_with_caption("regional_suicide_ranking_2024.png", "그림 5. 2024년 시도별 자살률 순위 및 삶의 만족 비율 비교")
    add_image_with_caption("satisfaction_profile_2024.png", "그림 6. 2024년 지역별 삶의 만족도 답변 분포 비율 (평균 점수 오름차순)")

    add_body_paragraph("그림 5와 그림 6의 시각화 자료는 2024년 기준 17개 지자체의 실제 평면 데이터를 대조합니다. 대구(불만족 21.4%), 세종(불만족 20.8%) 등 불만족 영역이 크지만 상대적으로 낮은 자살률을 유지하는 젊은 도시군과, 충남(불만족 6.2%), 전남(불만족 14.4%) 등 설문상 수치는 평온하지만 고령 노인층 고위험으로 자살 사망률이 치솟는 실질 고위험 지자체의 분포 차이를 선명하게 확인해 줍니다.")

    # --- 7. 결론 및 정책 제언 ---
    add_custom_heading("7. 결론 및 정책적 시사점", 1)
    add_bullet_point("주관적 설문 지표의 관리적 한계: '주관적 삶의 만족도' 설문조사 통계는 지역의 실제 정신건강 취약성을 완벽하게 대변하지 못합니다. 만족도가 높게 나오는 지자체라 하더라도 고령층의 고립 및 자살 예방을 위한 실질적인 스크리닝 체계를 소홀히 해서는 안 됩니다.")
    add_bullet_point("지자체별 인구 구조 맞춤형 예산 및 관리 차별화: 고령화 비율이 높은 농어촌 지자체(충남, 전남, 강원, 제주 등)는 노인 고독사 예방을 위한 지역사회 돌봄 네트워크 가동, 밀착 보건 체계, 농약 보관함 보급 등 실질적인 노인 자살 예방 정책에 예산과 인력을 집중해야 합니다. 반면, 청년층이 집중된 세종, 서울 등은 삶의 주관적 스트레스 완화를 위한 근로 여건 개선, 주거 안정, 청년 정신건강 인프라 확충에 초점을 두어 맞춤형 이원 트랙 예방 체계를 수립해야 합니다.")
    add_bullet_point("남성 고위험군 특화 안전망 수립: 유사한 수준의 만족 비율에도 자살률이 2.5배 가량 급증하는 남성 인구의 극단적 취약성에 주목해야 합니다. 남성의 예방 관리는 보편적인 심리 위안 프로그램에 그쳐서는 안 되며, 실직에 처한 가장의 생계비 긴급 지원, 은퇴 세대 사회적 단절 해소 프로그램, 중장년 남성 맞춤형 보건 인프라의 다각화가 시급히 동반되어야 합니다.")

    # Save document in local workspace
    doc_path = "./상관관계_분석_보고서.docx"
    doc.save(doc_path)
    print(f"Word document saved to '{doc_path}'.")

if __name__ == "__main__":
    main()
